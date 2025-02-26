import streamlit as st
from streamlit.runtime.uploaded_file_manager import UploadedFile
from docx import Document
from docx.document import Document as DocumentObject
from docx.shared import Mm
from io import BytesIO
from PIL import Image
import re
from logging import getLogger

logging = getLogger(__name__)

# Max image size (200 MB)
MAX_IMAGE_SIZE = 200 * 1024 * 1024

def count_existing_images(doc: DocumentObject, prefix: str) -> int:
    """
    Count images already appended by scanning paragraphs for labels 
    that start with the given prefix, expecting the format prefix + number.
    """
    max_index = 0
    pattern = re.compile(rf"^{re.escape(prefix)}(\d+)$")
    for para in doc.paragraphs:
        match = pattern.match(para.text.strip())
        if match:
            try:
                num = int(match.group(1))
                if num > max_index:
                    max_index = num
            except ValueError:
                continue
    return max_index

def process_doc_upload() -> str:
    """
    Process a file and return its name.
    """
    uploaded_doc: UploadedFile = st.session_state['uploaded_doc']
    if uploaded_doc is not None:
        try:
            doc = Document(uploaded_doc)
            st.session_state['doc'] = doc
            st.session_state['doc_name'] = uploaded_doc.name
            st.session_state.pop('doc_io', None)  # Clear any previous download
            st.toast(f"Loaded document: {uploaded_doc.name}")
            return uploaded_doc.name
        except Exception as e:
            st.error(f"Error loading document: {e}")
            return None
    else:
        st.session_state.pop('doc', None)
        st.session_state.pop('doc_name', None)
        st.session_state.pop('doc_io', None)
        return None
    

def main():
    # Configure the page to use a wide layout so the sidebar does not overlap the main content
    st.set_page_config(page_title="Word Document Image Appender", layout="wide")

    st.title("Word Document Image Appender")

    # ---------------------------------
    # MAIN AREA: Document & Image Management
    # ---------------------------------
    st.subheader("Step 1: Create or Upload Document")

    doc_name = st.text_input("Document Name", value="New_Document")

    if st.button("Create New Document"):
        st.session_state['doc'] = Document()
        st.session_state['doc_name'] = doc_name + ".docx"
        st.session_state.pop('doc_io', None)  # Clear any previous download
        st.toast("New document created!")

    uploaded_doc = st.file_uploader("Upload Existing Document", type=["docx"], key='uploaded_doc', on_change=process_doc_upload)

    if 'doc_name' in st.session_state:
        st.info(f"**Current Document:** {st.session_state['doc_name']}")

    # Image Naming Mode
    st.subheader("Image Settings")
    naming_mode = st.radio(
        "Choose how you want to name images:",
        options=["Auto-numbering", "Custom naming"],
        index=0,
        key="naming_mode"
    )

    # If auto-numbering, allow user to set a prefix
    if naming_mode == "Auto-numbering":
        image_prefix = st.text_input("Image Prefix", value="Image", key="image_prefix")
        image_index = st.text_input("Image Index Start", value="1", key="image_index")
    else:
        st.info("You will be able to set a custom name for each image below.")
    
    image_width_mm = st.slider(
        "Image Width (millimeters)",
        min_value=50,
        max_value=300,
        value=150,
        step=10,
        help="Select the width for inserted images in millimeters"
    )


    # Image Uploader & Preview
    st.subheader("Select Images")
    uploaded_images = st.file_uploader(
        "Drag and drop or browse files",
        type=["jpg", "jpeg", "png", "gif"],
        accept_multiple_files=True
    )

    # Store the uploaded images in session_state so we can use them in the sidebar
    st.session_state['uploaded_images'] = uploaded_images if uploaded_images else []

    custom_names = []
    if uploaded_images:
        st.markdown("### Preview & Naming")
        for idx, image_file in enumerate(uploaded_images):
            # Check size
            if image_file.size > MAX_IMAGE_SIZE:
                st.error(f"**{image_file.name}** exceeds the maximum file size limit.")
                custom_names.append("")
                continue

            # Preview
            try:
                image_file.seek(0)
                img = Image.open(image_file)
                st.image(img, caption=image_file.name, width=200)
            except Exception as e:
                st.error(f"Error previewing {image_file.name}: {e}")
                custom_names.append("")
                continue

            # If custom naming, let user define the name
            if naming_mode == "Custom naming":
                base_name = image_file.name.rsplit(".", 1)[0]
                name = st.text_input(
                    f"Label for {image_file.name}",
                    value=base_name,
                    key=f"custom_name_{idx}"
                )
                custom_names.append(name)
            else:
                # For auto-numbering, we just push an empty placeholder
                custom_names.append("")

    # Keep custom names in session state
    st.session_state['custom_names'] = custom_names

    # ---------------------------------
    # SIDEBAR: Document Actions & Buttons
    # ---------------------------------
    with st.sidebar:
        st.header("Document Actions")

        # 3. Append Images Button (shown only if we have images & a loaded doc)
        if 'doc' in st.session_state and 'uploaded_images' in st.session_state:
            if st.session_state['uploaded_images']:
                if st.button("Append Images"):
                    doc: DocumentObject = st.session_state['doc']
                    naming_mode = st.session_state.get('naming_mode', 'Auto-numbering')
                    image_prefix = st.session_state.get('image_prefix', 'Image')
                    custom_names = st.session_state.get('custom_names', [])

                    # For auto-numbering, count existing images for next index
                    if naming_mode == "Auto-numbering":
                        current_index = count_existing_images(doc, image_prefix)

                    # Append each image
                    for idx, image_file in enumerate(st.session_state['uploaded_images']):
                        if image_file.size > MAX_IMAGE_SIZE:
                            st.error(f"**{image_file.name}** exceeds the maximum file size limit. Skipping.")
                            continue
                        try:
                            image_file.seek(0)
                            # Add spacing before the image
                            doc.add_paragraph("")
                            doc.add_picture(image_file, width=Mm(image_width_mm))

                            # Determine label
                            if naming_mode == "Auto-numbering":
                                label = f"{image_prefix}{current_index + int(image_index)}"
                                current_index += 1
                            else:
                                # Custom name (fallback to base filename if blank)
                                label = custom_names[idx] or image_file.name.rsplit(".", 1)[0]

                            # Add label below image
                            doc.add_paragraph(label)
                            # Extra spacing
                            doc.add_paragraph("")
                        except Exception as e:
                            st.error(f"Error appending {image_file.name}: {e}")

                    # Save to BytesIO
                    doc_io = BytesIO()
                    try:
                        doc.save(doc_io)
                        doc_io.seek(0)
                        st.session_state['doc_io'] = doc_io
                        st.success("Images appended successfully!")
                    except Exception as e:
                        st.error(f"Error saving document: {e}")
            else:
                st.button("Append Images (Upload Images First)", disabled=True)
        else:
            st.button("Append Images (Create or Upload Existing Document First)", disabled=True)

        # 4. Download Updated Document
        if 'doc_io' in st.session_state:
            st.download_button(
                label="Download Updated Document",
                data=st.session_state['doc_io'],
                file_name=st.session_state.get('doc_name', 'document.docx'),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.button("Download Updated Document (Append Images First)", disabled=True)

if __name__ == "__main__":
    main()
